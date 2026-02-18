import os
import hashlib
import json
import threading
import struct
from datetime import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from PIL import Image, ExifTags, ImageTk
import PyPDF2
import docx
import openpyxl
import exifread
import numpy as np


class ModernForensicsTool:
    def __init__(self, root):
        self.root = root
        self.root.title("🔐 HideU Metadata") 
        self.root.geometry("1400x850")
        
        self.center_window()
        
        # Color Scheme 
        self.colors = {
            'bg': '#0f172a',          
            'fg': '#e2e8f0',         
            'primary': '#38bdf8',     
            'secondary': '#1e293b',   
            'accent': '#818cf8',      
            'success': '#4ade80',    
            'warning': '#facc15',     
            'danger': '#f87171',     
            'card': '#1e293b',        
            'border': '#334155',     
            'select': '#0ea5e9'      
        }
        
        self.selected_file = tk.StringVar()
        self.selected_folder = tk.StringVar()
        self.secret_message = tk.StringVar()
        self.current_file_type = tk.StringVar(value="📁 WAITING FOR INPUT...")
        
        self.current_metadata = {}
        self.batch_results = []
        
        self.setup_styles()
        self.setup_gui()
        self.processing = False
        
        # Trace file selection to update stego preview
        self.selected_file.trace_add('write', self.on_file_selected)

    def center_window(self):
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')

    def setup_styles(self):
        style = ttk.Style()
        style.theme_use('clam')  
        
        # General configurations
        style.configure('.', background=self.colors['bg'], foreground=self.colors['fg'], font=('Segoe UI', 10))
        style.configure('TFrame', background=self.colors['bg'])
        
        # 1. Main Headers
        style.configure('Title.TLabel', background=self.colors['bg'], foreground=self.colors['primary'],
                       font=('Segoe UI', 26, 'bold'))
        style.configure('Subtitle.TLabel', background=self.colors['bg'], foreground=self.colors['fg'],
                       font=('Segoe UI', 11))
        
        # 2. Cards (Panels)
        style.configure('Card.TFrame', background=self.colors['card'], relief='flat')
        style.configure('CardTitle.TLabel', background=self.colors['card'], foreground=self.colors['primary'],
                       font=('Segoe UI', 12, 'bold'))
        style.configure('CardText.TLabel', background=self.colors['card'], foreground=self.colors['fg'],
                       font=('Segoe UI', 10))
        style.configure('Card.TLabelframe', background=self.colors['card'], foreground=self.colors['primary'],
                       relief='solid', borderwidth=1)
        style.configure('Card.TLabelframe.Label', background=self.colors['card'], foreground=self.colors['primary'],
                       font=('Segoe UI', 11, 'bold'))

        # 3. Buttons (Flat & Modern)
        # Primary
        style.configure('Primary.TButton', background=self.colors['primary'], foreground='#0f172a',
                       font=('Segoe UI', 10, 'bold'), borderwidth=0, focuscolor=self.colors['primary'])
        style.map('Primary.TButton', background=[('active', '#7dd3fc'), ('pressed', '#0284c7')]) 
        
        # Accent
        style.configure('Accent.TButton', background=self.colors['accent'], foreground='#0f172a',
                       font=('Segoe UI', 10, 'bold'), borderwidth=0, focuscolor=self.colors['accent'])
        style.map('Accent.TButton', background=[('active', '#a5b4fc')])
        
        # Success
        style.configure('Success.TButton', background=self.colors['success'], foreground='#0f172a',
                       font=('Segoe UI', 10, 'bold'), borderwidth=0)
        style.map('Success.TButton', background=[('active', '#86efac')])

        # 4. Inputs
        style.configure('Modern.TEntry', fieldbackground=self.colors['secondary'], 
                       foreground=self.colors['fg'], insertcolor='white', borderwidth=0)
        
        # 5. Notebook (Tabs)
        style.configure('Modern.TNotebook', background=self.colors['bg'], borderwidth=0)
        style.configure('Modern.TNotebook.Tab', background=self.colors['secondary'], foreground=self.colors['fg'],
                       padding=[20, 10], font=('Segoe UI', 10, 'bold'), borderwidth=0)
        style.map('Modern.TNotebook.Tab', 
                 background=[('selected', self.colors['primary']), ('active', self.colors['border'])],
                 foreground=[('selected', '#0f172a'), ('active', self.colors['fg'])])

        # 6. Treeview (Tables)
        style.configure('Modern.Treeview', background=self.colors['secondary'], 
                       foreground=self.colors['fg'], fieldbackground=self.colors['secondary'],
                       font=('Consolas', 10), rowheight=30, borderwidth=0)
        style.configure('Modern.Treeview.Heading', background=self.colors['border'], 
                       foreground=self.colors['primary'], font=('Segoe UI', 10, 'bold'), relief='flat')
        style.map('Modern.Treeview', background=[('selected', self.colors['select'])], 
                 foreground=[('selected', 'white')])

        # 7. Scrollbars
        style.configure('Modern.Vertical.TScrollbar', background=self.colors['secondary'],
                       troughcolor=self.colors['bg'], bordercolor=self.colors['bg'], arrowcolor=self.colors['primary'])

    def setup_gui(self):
        self.root.configure(bg=self.colors['bg'])
        
        # Main Padding Container
        main_container = ttk.Frame(self.root, style='TFrame')
        main_container.pack(fill='both', expand=True, padx=20, pady=20)

        # --- HEADER ---
        header_frame = ttk.Frame(main_container, style='TFrame')
        header_frame.pack(fill='x', pady=(0, 25))
        
        # Title Left
        title_box = ttk.Frame(header_frame, style='TFrame')
        title_box.pack(side='left')
        # --- NAME CHANGE HERE ---
        ttk.Label(title_box, text="🔐 HideU Metadata", style='Title.TLabel').pack(anchor='w')
        ttk.Label(title_box, text="DIGITAL FORENSICS & STEGANOGRAPHY SUITE", style='Subtitle.TLabel').pack(anchor='w')

        # File Type Indicator Right
        type_box = ttk.Frame(header_frame, style='Card.TFrame', padding=15)
        type_box.pack(side='right')
        ttk.Label(type_box, text="DETECTED TYPE", style='CardText.TLabel', font=('Segoe UI', 8, 'bold')).pack(anchor='e')
        ttk.Label(type_box, textvariable=self.current_file_type, style='CardTitle.TLabel', font=('Segoe UI', 14)).pack(anchor='e')

        # --- BODY ---
        content_frame = ttk.Frame(main_container, style='TFrame')
        content_frame.pack(fill='both', expand=True)

        # LEFT SIDEBAR
        sidebar = ttk.Frame(content_frame, style='TFrame', width=320)
        sidebar.pack(side='left', fill='y', padx=(0, 20))
        sidebar.pack_propagate(False)

        # RIGHT RESULTS
        results_panel = ttk.Frame(content_frame, style='TFrame')
        results_panel.pack(side='right', fill='both', expand=True)

        # Build UI Components
        self.build_sidebar(sidebar)

        # Tabs
        self.notebook = ttk.Notebook(results_panel, style='Modern.TNotebook')
        self.notebook.pack(fill='both', expand=True)

        # Tab 1: Metadata
        meta_tab = ttk.Frame(self.notebook, style='TFrame')
        self.notebook.add(meta_tab, text='📊 METADATA')
        self.setup_metadata_tab(meta_tab)

        # Tab 2: Steganography
        self.stego_tab = ttk.Frame(self.notebook, style='TFrame')
        self.notebook.add(self.stego_tab, text='🔓 STEGO LAB')
        self.setup_stego_tab(self.stego_tab)

        # Tab 3: Batch
        batch_tab = ttk.Frame(self.notebook, style='TFrame')
        self.notebook.add(batch_tab, text='📦 BATCH SCAN')
        self.setup_batch_tab(batch_tab)

        # Tab 4: Hex
        hex_tab = ttk.Frame(self.notebook, style='TFrame')
        self.notebook.add(hex_tab, text='🔢 HEX VIEW')
        self.setup_hex_tab(hex_tab)

        # --- STATUS BAR ---
        status_frame = ttk.Frame(main_container, style='Card.TFrame', padding=10)
        status_frame.pack(fill='x', pady=(20, 0))
        
        self.status_label = ttk.Label(status_frame, text="✅ SYSTEM READY", style='CardText.TLabel')
        self.status_label.pack(side='left')
        
        self.progress = ttk.Progressbar(status_frame, mode='indeterminate', style='Horizontal.TProgressbar')
        self.progress.pack(side='right', fill='x', expand=True, padx=(20, 0))

    def build_sidebar(self, parent):
        # File Selection Card
        file_card = ttk.LabelFrame(parent, text=" SOURCE SELECTION ", style='Card.TLabelframe', padding=15)
        file_card.pack(fill='x', pady=(0, 15))

        ttk.Label(file_card, text="Target File Path:", style='CardText.TLabel').pack(anchor='w', pady=(0,5))
        
        # File Entry & Button
        f_entry = ttk.Frame(file_card, style='Card.TFrame')
        f_entry.pack(fill='x', pady=(0, 10))
        ttk.Entry(f_entry, textvariable=self.selected_file, style='Modern.TEntry').pack(side='left', fill='x', expand=True, padx=(0,5))
        ttk.Button(f_entry, text="📂", width=3, command=self.browse_file, style='Primary.TButton').pack(side='right')

        ttk.Button(file_card, text="📂 Select Folder (Batch)", command=self.browse_folder, style='Accent.TButton').pack(fill='x', pady=(0,10))
        
        # Action Button
        ttk.Button(file_card, text="⚡ EXTRACT METADATA", command=self.extract_single_file, style='Primary.TButton').pack(fill='x')

        # Stego Card
        stego_card = ttk.LabelFrame(parent, text=" STEGO OPERATIONS ", style='Card.TLabelframe', padding=15)
        stego_card.pack(fill='x', pady=(0, 15))

        ttk.Label(stego_card, text="Secret Message / Key:", style='CardText.TLabel').pack(anchor='w', pady=(0,5))
        ttk.Entry(stego_card, textvariable=self.secret_message, style='Modern.TEntry').pack(fill='x', pady=(0,10))

        row = ttk.Frame(stego_card, style='Card.TFrame')
        row.pack(fill='x', pady=(0,5))
        ttk.Button(row, text="🔒 HIDE", command=self.encode_steganography, style='Primary.TButton').pack(side='left', fill='x', expand=True, padx=(0,2))
        ttk.Button(row, text="🔓 REVEAL", command=self.decode_steganography, style='Success.TButton').pack(side='right', fill='x', expand=True, padx=(2,0))
        
        ttk.Button(stego_card, text="🔎 ANALYZE IMAGE", command=self.analyze_steganography, style='Accent.TButton').pack(fill='x', pady=(5,0))

        # Export Card
        export_card = ttk.LabelFrame(parent, text=" REPORTING ", style='Card.TLabelframe', padding=15)
        export_card.pack(fill='x')
        ttk.Button(export_card, text="💾 EXPORT JSON", command=self.export_to_json, style='Primary.TButton').pack(fill='x', pady=(0,5))
        ttk.Button(export_card, text="🗑️ CLEAR SESSION", command=self.clear_results, style='Accent.TButton').pack(fill='x')

    # ---------- Metadata Tab ----------
    def setup_metadata_tab(self, parent):
        # Stats Top Bar
        stats_frame = ttk.Frame(parent, style='TFrame')
        stats_frame.pack(fill='x', pady=15, padx=5)
        
        self.stats_labels = {}
        icons = {'Files': '📄', 'Size': '💾', 'Hashes': '🔑', 'Dates': '📅'}
        
        for stat, icon in icons.items():
            card = ttk.Frame(stats_frame, style='Card.TFrame', padding=10)
            card.pack(side='left', fill='x', expand=True, padx=5)
            
            ttk.Label(card, text=f"{icon} {stat}", style='CardText.TLabel', font=('Segoe UI', 9, 'bold')).pack()
            self.stats_labels[stat] = ttk.Label(card, text="-", style='CardTitle.TLabel', font=('Segoe UI', 12))
            self.stats_labels[stat].pack()

        # Treeview
        tree_frame = ttk.Frame(parent, style='TFrame')
        tree_frame.pack(fill='both', expand=True, padx=5, pady=5)

        self.tree = ttk.Treeview(tree_frame, columns=('Property', 'Value'), show='tree',
                                style='Modern.Treeview', height=20)
        
        # Scrollbars
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview, style='Modern.Vertical.TScrollbar')
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview, style='Modern.Horizontal.TScrollbar')
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.heading('#0', text='CATEGORY', anchor='w')
        self.tree.heading('Property', text='PROPERTY', anchor='w')
        self.tree.heading('Value', text='VALUE', anchor='w')
        
        self.tree.column('#0', width=250)
        self.tree.column('Property', width=250)
        self.tree.column('Value', width=600)

        self.tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')
        
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)

    # ---------- Steganography Tab ----------
    def setup_stego_tab(self, parent):
        paned = ttk.PanedWindow(parent, orient='horizontal')
        paned.pack(fill='both', expand=True, padx=10, pady=10)

        # LEFT: Image Preview
        left_panel = ttk.LabelFrame(paned, text=" VISUAL PREVIEW ", style='Card.TLabelframe', padding=10)
        paned.add(left_panel, weight=1)

        self.preview_label = ttk.Label(left_panel, style='CardText.TLabel', anchor='center')
        self.preview_label.pack(fill='both', expand=True, pady=10)
        
        info_frame = ttk.Frame(left_panel, style='Card.TFrame')
        info_frame.pack(fill='x')
        
        self.img_format_label = ttk.Label(info_frame, text="FORMAT: --", style='CardText.TLabel', font=('Consolas', 9))
        self.img_format_label.pack(anchor='w')
        self.img_dimensions_label = ttk.Label(info_frame, text="SIZE: --", style='CardText.TLabel', font=('Consolas', 9))
        self.img_dimensions_label.pack(anchor='w')
        self.img_capacity_label = ttk.Label(info_frame, text="CAPACITY: --", style='CardText.TLabel', foreground=self.colors['success'], font=('Consolas', 9))
        self.img_capacity_label.pack(anchor='w')

        # RIGHT: Text Log
        right_panel = ttk.LabelFrame(paned, text=" ANALYSIS LOG ", style='Card.TLabelframe', padding=10)
        paned.add(right_panel, weight=2)

        self.stego_text = scrolledtext.ScrolledText(right_panel, wrap=tk.WORD,
                                                   bg=self.colors['secondary'], fg=self.colors['fg'],
                                                   insertbackground='white',
                                                   font=('Consolas', 10), borderwidth=0, padx=10, pady=10)
        self.stego_text.pack(fill='both', expand=True)

    # ---------- Batch Tab ----------
    def setup_batch_tab(self, parent):
        # Controls
        controls = ttk.Frame(parent, style='TFrame')
        controls.pack(fill='x', padx=10, pady=10)
        
        ttk.Button(controls, text="▶ START BATCH", command=self.process_batch, style='Success.TButton').pack(side='left', padx=5)
        ttk.Button(controls, text="📊 SAVE REPORT", command=self.generate_batch_report, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(controls, text="🗑️ CLEAR", command=lambda: self.batch_tree.delete(*self.batch_tree.get_children()), style='Accent.TButton').pack(side='right', padx=5)

        # Tree
        tree_frame = ttk.Frame(parent, style='TFrame')
        tree_frame.pack(fill='both', expand=True, padx=10, pady=(0,10))

        cols = ('File', 'Type', 'Size', 'MD5', 'Status', 'Modified')
        self.batch_tree = ttk.Treeview(tree_frame, columns=cols, show='headings', style='Modern.Treeview')
        
        for col in cols:
            self.batch_tree.heading(col, text=col.upper())
            self.batch_tree.column(col, width=100)

        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.batch_tree.yview, style='Modern.Vertical.TScrollbar')
        self.batch_tree.configure(yscrollcommand=vsb.set)
        
        self.batch_tree.pack(side='left', fill='both', expand=True)
        vsb.pack(side='right', fill='y')

    # ---------- Hex Tab ----------
    def setup_hex_tab(self, parent):
        top_bar = ttk.Frame(parent, style='TFrame')
        top_bar.pack(fill='x', padx=10, pady=10)
        
        ttk.Button(top_bar, text="LOAD HEX FOR CURRENT FILE", command=self.load_hex_view, style='Primary.TButton').pack(anchor='w')

        hex_frame = ttk.Frame(parent, style='TFrame')
        hex_frame.pack(fill='both', expand=True, padx=10, pady=(0,10))

        self.hex_text = scrolledtext.ScrolledText(hex_frame, wrap=tk.NONE,
                                                 bg='#0f172a', fg='#38bdf8',  # Dark BG, Cyan Text
                                                 insertbackground='white',
                                                 font=('Consolas', 10), borderwidth=0)
        self.hex_text.pack(fill='both', expand=True)

    # ---------- Utility Methods ----------
    def update_status(self, message, is_error=False):
        color = self.colors['danger'] if is_error else self.colors['success']
        if not is_error and "..." in message: color = self.colors['fg']
            
        self.status_label.configure(text=message, foreground=color)
        self.root.update()

    def show_progress(self, show=True):
        if show:
            self.progress.start(10)
        else:
            self.progress.stop()

    def browse_file(self):
        f = filedialog.askopenfilename(title="Select File for Analysis")
        if f:
            self.selected_file.set(f)
            self.update_file_type(f)
            self.update_status(f"✅ Selected: {os.path.basename(f)}")

    def update_file_type(self, filename):
        ext = os.path.splitext(filename)[1].lower()
        icons = {'.jpg':'🖼️','.jpeg':'🖼️','.png':'🖼️','.gif':'🖼️','.bmp':'🖼️','.tiff':'🖼️',
                '.pdf':'📄','.docx':'📝','.doc':'📝','.xlsx':'📊','.txt':'📃','.zip':'📦',
                '.rar':'📦','.7z':'📦','.exe':'⚙️','.dll':'⚙️'}
        icon = icons.get(ext, '📁')
        self.current_file_type.set(f"{icon} {ext[1:].upper() if ext else 'UNKNOWN'}")

    def browse_folder(self):
        f = filedialog.askdirectory(title="Select Folder for Batch Processing")
        if f:
            self.selected_folder.set(f)
            self.update_status(f"✅ Selected folder: {os.path.basename(f)}")

    # ---------- Stego Preview (triggered on file selection) ----------
    def on_file_selected(self, *args):
        self.update_stego_preview()

    def update_stego_preview(self):
        """Update the preview panel with the currently selected image (if any)."""
        path = self.selected_file.get()
        if not path or not os.path.exists(path):
            self.preview_label.config(image='', text='NO IMAGE SELECTED')
            self.img_format_label.config(text="FORMAT: --")
            self.img_dimensions_label.config(text="SIZE: --")
            self.img_capacity_label.config(text="CAPACITY: --")
            return

        try:
            img = Image.open(path)
            if img.mode not in ('RGB', 'L'):
                img = img.convert('RGB')
            
            # Create thumbnail (max 300x300 for the new UI)
            img_copy = img.copy()
            img_copy.thumbnail((300, 300))
            photo = ImageTk.PhotoImage(img_copy)
            self.preview_label.config(image=photo, text='')
            self.preview_label.image = photo 

            ext = os.path.splitext(path)[1].upper()[1:] or 'UNKNOWN'
            self.img_format_label.config(text=f"FORMAT: {ext}")
            self.img_dimensions_label.config(text=f"SIZE: {img.width} × {img.height}")

            if img.mode == 'RGB':
                pixels = img.width * img.height
                capacity_bits = pixels * 3
            else:
                pixels = img.width * img.height
                capacity_bits = pixels * 3

            capacity_bytes = capacity_bits // 8
            self.img_capacity_label.config(text=f"CAPACITY: ~{capacity_bytes:,} BYTES")

        except Exception as e:
            self.preview_label.config(image='', text=f'PREVIEW ERROR')

    # ---------- METADATA EXTRACTION ----------
    def extract_single_file(self):
        if not self.selected_file.get():
            messagebox.showwarning("No File", "Please select a file first.", parent=self.root)
            return
        threading.Thread(target=self._extract_metadata_thread, daemon=True).start()

    def _extract_metadata_thread(self):
        self.show_progress(True)
        self.update_status("📊 Extracting metadata...")
        try:
            path = self.selected_file.get()
            metadata = self.extract_all_metadata(path)
            self.root.after(0, lambda: self.populate_metadata(metadata))
            self.current_metadata = metadata
        except Exception as e:
            self.root.after(0, lambda: self.update_status(f"❌ Error: {str(e)}", is_error=True))
        finally:
            self.root.after(0, lambda: self.show_progress(False))

    def extract_all_metadata(self, path):
        meta = {}
        stat = os.stat(path)
        meta['Basic Info'] = {
            'File Name': os.path.basename(path),
            'File Size': f"{stat.st_size:,} bytes",
            'Created': datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
            'Modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
            'Accessed': datetime.fromtimestamp(stat.st_atime).strftime('%Y-%m-%d %H:%M:%S'),
            'Extension': os.path.splitext(path)[1],
        }

        meta['Hashes'] = self.compute_hashes(path)

        ext = os.path.splitext(path)[1].lower()
        if ext in ('.jpg','.jpeg','.png','.gif','.bmp','.tiff'):
            meta.update(self.extract_image_metadata(path))
        elif ext == '.pdf':
            meta.update(self.extract_pdf_metadata(path))
        elif ext == '.docx':
            meta.update(self.extract_docx_metadata(path))
        elif ext == '.xlsx':
            meta.update(self.extract_xlsx_metadata(path))

        return meta

    def compute_hashes(self, path, blocksize=65536):
        hashes = {}
        md5 = hashlib.md5()
        sha1 = hashlib.sha1()
        sha256 = hashlib.sha256()
        with open(path, 'rb') as f:
            while chunk := f.read(blocksize):
                md5.update(chunk)
                sha1.update(chunk)
                sha256.update(chunk)
        hashes['MD5'] = md5.hexdigest()
        hashes['SHA-1'] = sha1.hexdigest()
        hashes['SHA-256'] = sha256.hexdigest()
        return hashes

    def extract_image_metadata(self, path):
        meta = {}
        try:
            img = Image.open(path)
            info = {
                'Format': img.format,
                'Mode': img.mode,
                'Width': img.width,
                'Height': img.height,
                'Is Animated': getattr(img, 'is_animated', False),
                 'Frames': getattr(img, 'n_frames', 1)
            }
            meta['Image Properties'] = info

            # EXIF
            exif_data = {}
            if hasattr(img, '_getexif') and img._getexif():
                exif = img._getexif()
                for tag_id, value in exif.items():
                    tag = ExifTags.TAGS.get(tag_id, tag_id)
                    exif_data[tag] = str(value)
            if exif_data:
                meta['EXIF'] = exif_data

            # Use exifread for more detailed EXIF
            with open(path, 'rb') as f:
                tags = exifread.process_file(f, details=False)
                if tags:
                    gps = {}
                    for k, v in tags.items():
                        if 'GPS' in k:
                            gps[k] = str(v)
                    if gps:
                        meta['GPS'] = gps
        except Exception as e:
            meta['Image Error'] = str(e)
        return meta

    def extract_pdf_metadata(self, path):
        meta = {}
        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                info = reader.metadata
                if info:
                    meta['PDF Metadata'] = {k: str(v) for k, v in info.items()}
                meta['PDF Info'] = {
                    'Number of Pages': len(reader.pages),
                    'Encrypted': reader.is_encrypted
                }
        except Exception as e:
            meta['PDF Error'] = str(e)
        return meta

    def extract_docx_metadata(self, path):
        meta = {}
        try:
            doc = docx.Document(path)
            core_props = doc.core_properties
            props = {
                'Author': core_props.author,
                'Title': core_props.title,
                'Subject': core_props.subject,
                'Keywords': core_props.keywords,
                'Comments': core_props.comments,
                'Category': core_props.category,
                'Created': core_props.created,
                'Modified': core_props.modified,
                'Last Modified By': core_props.last_modified_by,
                'Revision': core_props.revision,
                'Word Count': len(doc.paragraphs) + len(doc.tables)
            }
            meta['DOCX Properties'] = {k: str(v) for k, v in props.items() if v}
        except Exception as e:
            meta['DOCX Error'] = str(e)
        return meta

    def extract_xlsx_metadata(self, path):
        meta = {}
        try:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            props = wb.properties
            p = {
                'Title': props.title,
                'Subject': props.subject,
                'Creator': props.creator,
                'Keywords': props.keywords,
                'Description': props.description,
                'Created': props.created,
                'Modified': props.modified,
                'Last Modified By': props.lastModifiedBy,
                'Category': props.category,
                'Sheets': wb.sheetnames
            }
            meta['XLSX Properties'] = {k: str(v) for k, v in p.items() if v}
        except Exception as e:
            meta['XLSX Error'] = str(e)
        return meta

    def populate_metadata(self, metadata):
        self.clear_tree()
        for cat, data in metadata.items():
            parent = self.tree.insert('', 'end', text=cat, open=True)
            if isinstance(data, dict):
                for k, v in data.items():
                    self.tree.insert(parent, 'end', values=(k, v))
            else:
                self.tree.insert(parent, 'end', values=(cat, str(data)))
        
        # --- FIX: Pass correct metadata to stats updater ---
        self.update_stats(metadata)
        self.update_status("✅ Metadata extracted successfully!")
        self.notebook.select(0)

    def update_stats(self, metadata):
        # 1. Update Files
        self.stats_labels['Files'].configure(text="1")

        # 2. Update Size
        size = metadata.get('Basic Info', {}).get('File Size', '-')
        self.stats_labels['Size'].configure(text=size)

        # 3. Update Hashes
        hash_count = len(metadata.get('Hashes', {}))
        if hash_count > 0:
            self.stats_labels['Hashes'].configure(text=f"{hash_count} Generated")
        else:
            self.stats_labels['Hashes'].configure(text="-")

        # 4. Update Dates (Modified Date)
        full_date = metadata.get('Basic Info', {}).get('Modified', '-')
        if full_date != '-':
            display_date = full_date.split(' ')[0] # Get YYYY-MM-DD
        else:
            display_date = "-"
        self.stats_labels['Dates'].configure(text=display_date)

    def clear_tree(self):
        for i in self.tree.get_children():
            self.tree.delete(i)

    # ---------- STEGANOGRAPHY ----------
    def encode_steganography(self):
        if not self.selected_file.get():
            messagebox.showwarning("No File", "Please select an image file.", parent=self.root)
            return
        if not self.secret_message.get():
            messagebox.showwarning("No Message", "Enter a message to encode.", parent=self.root)
            return

        ext = os.path.splitext(self.selected_file.get())[1].lower()
        if ext in ('.jpg', '.jpeg'):
            if not messagebox.askyesno("Lossy Format Warning",
                                       "JPEG compression will corrupt the hidden message.\n"
                                       "Proceed anyway?",
                                       parent=self.root):
                return

        threading.Thread(target=self._encode_thread, daemon=True).start()

    def _encode_thread(self):
        self.show_progress(True)
        self.update_status("🔒 Encoding message...")
        try:
            path = self.selected_file.get()
            msg = self.secret_message.get()

            msg_bytes = msg.encode('utf-8')
            header = struct.pack('>I', len(msg_bytes))
            payload = header + msg_bytes
            bits = ''.join(format(b, '08b') for b in payload)

            img = Image.open(path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            arr = np.array(img)
            
            if arr.dtype != np.uint8:
                arr = arr.astype(np.uint8)
                
            flat = arr.flatten()

            if len(bits) > len(flat):
                raise ValueError(f"Image too small. Need {len(bits)} bits, have {len(flat)}.")

            for i, bit in enumerate(bits):
                flat[i] = (flat[i] & 0xFE) | int(bit)

            new_arr = flat.reshape(arr.shape)
            new_img = Image.fromarray(new_arr, img.mode)

            save_path = filedialog.asksaveasfilename(
                defaultextension=".png",
                filetypes=[("PNG files", "*.png")],
                title="Save encoded image as PNG"
            )
            if save_path:
                new_img.save(save_path, 'PNG')
                self.root.after(0, lambda: self.update_status(
                    f"✅ Saved to {os.path.basename(save_path)}"))
                self.root.after(0, lambda: self.stego_display_success(
                    "Encoding completed successfully."))
            else:
                self.root.after(0, lambda: self.update_status("Encoding cancelled."))

        except Exception as e:
            self.root.after(0, lambda: self.update_status(
                f"❌ Encoding failed: {str(e)}", is_error=True))
        finally:
            self.root.after(0, lambda: self.show_progress(False))

    def decode_steganography(self):
        if not self.selected_file.get():
            messagebox.showwarning("No File", "Please select an image file.", parent=self.root)
            return
        threading.Thread(target=self._decode_thread, daemon=True).start()

    def _decode_thread(self):
        self.show_progress(True)
        self.update_status("🔍 Decoding message...")
        try:
            path = self.selected_file.get()

            img = Image.open(path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            arr = np.array(img)
            if arr.dtype != np.uint8:
                arr = arr.astype(np.uint8)
            flat = arr.flatten()

            length_bits = ''.join(str(flat[i] & 1) for i in range(32))
            length = int(length_bits, 2)

            if length <= 0 or length > 10000:
                raise ValueError("No valid message header found.")

            total_bits = 32 + length * 8
            if total_bits > len(flat):
                raise ValueError("Image does not contain enough bits.")
            bits = ''.join(str(flat[i] & 1) for i in range(total_bits))

            bytes_data = bytes(
                int(bits[i:i+8], 2) for i in range(32, len(bits), 8)
            )

            msg = bytes_data.decode('utf-8', errors='replace')
            self.root.after(0, lambda: self.stego_display_message(msg))

        except Exception as e:
            self.root.after(0, lambda: self.update_status(
                f"❌ Decoding failed: {str(e)}", is_error=True))
        finally:
            self.root.after(0, lambda: self.show_progress(False))

    def stego_display_success(self, msg):
        self.stego_text.delete(1.0, tk.END)
        self.stego_text.insert(tk.END, f"✅ {msg}\n\n")
        self.notebook.select(1)

    def stego_display_message(self, msg):
        self.stego_text.delete(1.0, tk.END)
        self.stego_text.insert(tk.END, "🔍 EXTRACTED MESSAGE:\n\n")
        self.stego_text.insert(tk.END, msg)
        self.notebook.select(1)
        self.update_status("✅ Message decoded successfully!")

    def analyze_steganography(self):
        if not self.selected_file.get():
            messagebox.showwarning("No File", "Please select an image file.", parent=self.root)
            return
        threading.Thread(target=self._analyze_thread, daemon=True).start()

    def _analyze_thread(self):
        self.show_progress(True)
        self.update_status("🔬 Analyzing for steganography...")
        try:
            path = self.selected_file.get()
            img = Image.open(path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            arr = np.array(img)
            if arr.dtype != np.uint8:
                arr = arr.astype(np.uint8)
            flat = arr.flatten()
            total_pixels = len(flat)
            lsb_ones = int(np.sum(flat & 1))
            lsb_ratio = lsb_ones / total_pixels
            entropy = self._calculate_entropy(flat)
            stat = os.stat(path)
            file_size = stat.st_size
            expected_size = arr.shape[0] * arr.shape[1] * 3 + 1000
            size_ratio = file_size / expected_size if expected_size else 1

            report = []
            report.append("🔬 STEGANOGRAPHY ANALYSIS REPORT")
            report.append("=" * 60)
            report.append(f"File: {os.path.basename(path)}")
            report.append(f"Size: {file_size:,} bytes")
            report.append(f"Dimensions: {arr.shape[1]}×{arr.shape[0]}")
            report.append(f"Mode: {img.mode}")
            report.append("")
            report.append("📊 STATISTICAL ANALYSIS:")
            report.append(f"  LSB 1s: {lsb_ones:,} / {total_pixels:,} ({lsb_ratio:.2%})")
            report.append(f"  Entropy: {entropy:.4f} bits")
            report.append(f"  File size vs expected: {size_ratio:.2f}x")
            report.append("")
            report.append("⚠️  SUSPICION INDICATORS:")
            suspicious = []
            if abs(lsb_ratio - 0.5) > 0.1:
                suspicious.append("LSB distribution not uniform (potential stego)")
            if entropy > 7.9:
                suspicious.append("High entropy (possible encrypted/carved data)")
            if size_ratio > 1.2:
                suspicious.append("File significantly larger than expected")
            if suspicious:
                for s in suspicious:
                    report.append(f"  • {s}")
            else:
                report.append("  No obvious steganographic signs.")
            
            self.root.after(0, lambda: self.stego_display_report("\n".join(report)))
        except Exception as e:
            self.root.after(0, lambda: self.update_status(f"❌ Analysis failed: {str(e)}", is_error=True))
        finally:
            self.root.after(0, lambda: self.show_progress(False))

    def _calculate_entropy(self, data):
        values, counts = np.unique(data, return_counts=True)
        probs = counts / len(data)
        entropy = -np.sum(probs * np.log2(probs))
        return entropy

    def stego_display_report(self, report):
        self.stego_text.delete(1.0, tk.END)
        self.stego_text.insert(tk.END, report)
        self.notebook.select(1)
        self.update_status("✅ Analysis complete.")

    # ---------- BATCH PROCESSING ----------
    def process_batch(self):
        if not self.selected_folder.get():
            messagebox.showwarning("No Folder", "Please select a folder first.", parent=self.root)
            return
        threading.Thread(target=self._batch_thread, daemon=True).start()

    def _batch_thread(self):
        self.show_progress(True)
        self.update_status("📦 Processing batch...")
        self.root.after(0, lambda: self.batch_tree.delete(*self.batch_tree.get_children()))
        self.batch_results.clear()
        folder = self.selected_folder.get()
        try:
            files = []
            for root, dirs, files_ in os.walk(folder):
                for f in files_:
                    files.append(os.path.join(root, f))
                break  # top-level only
            total = len(files)
            for idx, path in enumerate(files):
                try:
                    stat = os.stat(path)
                    size = stat.st_size
                    ext = os.path.splitext(path)[1].lower()
                    modified = datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d')
                    md5 = self.compute_hashes(path)['MD5'][:8] + '...'
                    status = "✅ Clean"
                    if size > 10_000_000:
                        status = "⚠️ Large"
                    if ext in ('.exe', '.dll', '.bat', '.vbs', '.elf'):
                        status = "🔴 Executable"
                    self.root.after(0, lambda p=path, e=ext, s=size, m=md5, mod=modified, stat=status:
                                   self.batch_tree.insert('', 'end', values=(
                                       os.path.basename(p), e, f"{s:,}", m, stat, mod)))
                    self.batch_results.append({
                        'file': path, 'type': ext, 'size': size,
                        'md5': md5, 'status': status, 'modified': modified
                    })
                    self.root.after(0, lambda i=idx, t=total: self.update_status(f"📦 Processed {i+1}/{t} files..."))
                except Exception:
                    continue
            self.root.after(0, lambda: self.update_status(f"✅ Batch complete: {len(files)} files processed"))
        except Exception as e:
            self.root.after(0, lambda: self.update_status(f"❌ Batch error: {str(e)}", is_error=True))
        finally:
            self.root.after(0, lambda: self.show_progress(False))

    def generate_batch_report(self):
        if not self.batch_results:
            messagebox.showinfo("No Data", "No batch results to export.", parent=self.root)
            return
        report = "Batch Analysis Report\n"
        report += "="*50 + "\n"
        report += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        report += f"Folder: {self.selected_folder.get()}\n"
        report += f"Files analyzed: {len(self.batch_results)}\n\n"
        clean = sum(1 for r in self.batch_results if r['status'] == '✅ Clean')
        suspicious = len(self.batch_results) - clean
        report += f"Clean files: {clean}\n"
        report += f"Suspicious files: {suspicious}\n\n"
        report += "Details:\n"
        for r in self.batch_results:
            report += f"  {r['file']} - {r['type']} - {r['size']} bytes - {r['md5']} - {r['status']}\n"
        filename = filedialog.asksaveasfilename(defaultextension=".txt",
                                                filetypes=[("Text files","*.txt")],
                                                title="Save Batch Report")
        if filename:
            with open(filename, 'w') as f:
                f.write(report)
            self.update_status(f"✅ Report saved to {os.path.basename(filename)}")

    # ---------- HEX VIEWER ----------
    def load_hex_view(self):
        if not self.selected_file.get():
            messagebox.showwarning("No File", "Please select a file first.", parent=self.root)
            return
        threading.Thread(target=self._hex_thread, daemon=True).start()

    def _hex_thread(self):
        self.show_progress(True)
        self.update_status("🔢 Loading hex view...")
        try:
            path = self.selected_file.get()
            hex_dump = self.generate_hex_dump(path)
            self.root.after(0, lambda: self.display_hex(hex_dump))
        except Exception as e:
            self.root.after(0, lambda: self.update_status(f"❌ Hex error: {str(e)}", is_error=True))
        finally:
            self.root.after(0, lambda: self.show_progress(False))

    def generate_hex_dump(self, path, bytes_per_line=16):
        lines = []
        with open(path, 'rb') as f:
            offset = 0
            while chunk := f.read(bytes_per_line):
                hex_str = ' '.join(f'{b:02X}' for b in chunk)
                ascii_str = ''.join(chr(b) if 32 <= b <= 126 else '.' for b in chunk)
                lines.append(f'{offset:08X}: {hex_str:<{bytes_per_line*3}} {ascii_str}')
                offset += len(chunk)
                if offset > 2048:
                    lines.append('... (file truncated)')
                    break
        return '\n'.join(lines)

    def display_hex(self, hex_data):
        self.hex_text.delete(1.0, tk.END)
        self.hex_text.insert(tk.END, hex_data)
        self.update_status("✅ Hex view loaded successfully!")
        self.notebook.select(3)

    # ---------- EXPORT ----------
    def export_to_json(self):
        if not self.current_metadata and not self.batch_results:
            messagebox.showinfo("No Data", "No metadata or batch results to export.", parent=self.root)
            return
        data = {
            'timestamp': datetime.now().isoformat(),
            'metadata': self.current_metadata,
            'batch_results': self.batch_results
        }
        filename = filedialog.asksaveasfilename(defaultextension=".json",
                                                filetypes=[("JSON files","*.json")],
                                                title="Export to JSON")
        if filename:
            with open(filename, 'w') as f:
                json.dump(data, f, indent=2, default=str)
            self.update_status(f"✅ Exported to {os.path.basename(filename)}")

    def clear_results(self):
        self.clear_tree()
        self.batch_tree.delete(*self.batch_tree.get_children())
        self.stego_text.delete(1.0, tk.END)
        self.hex_text.delete(1.0, tk.END)
        self.current_metadata.clear()
        self.batch_results.clear()
        self.update_status("✅ All results cleared")


def main():
    root = tk.Tk()
    app = ModernForensicsTool(root)
    root.mainloop()


if __name__ == "__main__":
    main()